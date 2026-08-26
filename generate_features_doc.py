import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_document():
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    navy = RGBColor(30, 58, 138)       # #1e3a8a
    indigo = RGBColor(67, 56, 202)     # #4338ca
    slate_dark = RGBColor(30, 41, 59)  # #1e293b
    slate_gray = RGBColor(100, 116, 139) # #64748b

    # ----------------------------------------------------
    # COVER / HEADER TITLE
    # ----------------------------------------------------
    p_pre = doc.add_paragraph()
    r_pre = p_pre.add_run("ENTERPRISE SOFTWARE SPECIFICATION & ARCHITECTURE GUIDE")
    r_pre.font.size = Pt(9.5)
    r_pre.font.bold = True
    r_pre.font.color.rgb = indigo
    p_pre.paragraph_format.space_after = Pt(2)

    p_title = doc.add_paragraph()
    r_title = p_title.add_run("ProFlow: Enterprise Project Controls & Governance System")
    r_title.font.size = Pt(24)
    r_title.font.bold = True
    r_title.font.color.rgb = navy
    p_title.paragraph_format.space_after = Pt(4)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Complete Functional Catalog of Advanced Engineering Management, Financial Controls, APQP Stage-Gate Governance, and Real-Time Collaboration Engines")
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = slate_dark
    p_sub.paragraph_format.space_after = Pt(14)

    # Executive Metadata Table
    table_meta = doc.add_table(rows=5, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False

    meta_data = [
        ("Platform Name", "ProFlow (Autonomous Project Controls & APQP Governance Suite)"),
        ("Industry Focus", "Automotive, Aerospace, Defense, and Deep-Tech Hardware/Software"),
        ("System Architecture", "React 18 + Vite + TailwindCSS | Node.js Express | Prisma ORM | PostgreSQL"),
        ("Executive Leadership", "Parth Singh (Program Director & Chief EV Architect)"),
        ("Specification Version", "v2.5 Enterprise Production Release (2026 Edition)")
    ]

    for i, (label, val) in enumerate(meta_data):
        row = table_meta.rows[i]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.2)
        c1.width = Inches(4.6)
        set_cell_background(c0, "F1F5F9")
        set_cell_background(c1, "F8FAFC")
        set_cell_margins(c0, 120, 120, 150, 150)
        set_cell_margins(c1, 120, 120, 150, 150)

        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = slate_dark
        p0.paragraph_format.space_after = Pt(0)

        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(10)
        r1.font.color.rgb = slate_dark
        p1.paragraph_format.space_after = Pt(0)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ----------------------------------------------------
    # SECTION GENERATOR HELPER
    # ----------------------------------------------------
    def add_section_header(num, title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r_num = p.add_run(f"Module {num:02d} | ")
        r_num.font.bold = True
        r_num.font.size = Pt(11)
        r_num.font.color.rgb = indigo

        r_title = p.add_run(title)
        r_title.font.bold = True
        r_title.font.size = Pt(15)
        r_title.font.color.rgb = navy

    def add_bullet_feature(bold_title, desc):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r_b = p.add_run(bold_title + ": ")
        r_b.font.bold = True
        r_b.font.size = Pt(10)
        r_b.font.color.rgb = slate_dark

        r_d = p.add_run(desc)
        r_d.font.size = Pt(10)
        r_d.font.color.rgb = slate_dark

    # ----------------------------------------------------
    # MODULE 1: CPM & GANTT ENGINE
    # ----------------------------------------------------
    add_section_header(1, "Schedule & Critical Path Management (CPM / Gantt)")
    doc.add_paragraph("ProFlow features a mathematical scheduling engine that computes activity early/late dates, total float slack, and highlights critical path bottlenecks across complex multi-tier dependencies.")

    add_bullet_feature("CPM Forward & Backward Pass Algorithm", "Automatically calculates Early Start (ES), Early Finish (EF), Late Start (LS), Late Finish (LF), and Total Float (TF) for all network activities.")
    add_bullet_feature("Zero-Float Critical Path Highlighting", "Interactive toggle on the Gantt chart renders critical bottleneck tasks in high-visibility crimson red with animated indicators.")
    add_bullet_feature("4 Complete Dependency Types", "Full support for Finish-to-Start (FS), Start-to-Start (SS), Finish-to-Finish (FF), and Start-to-Finish (SF) dependency links.")
    add_bullet_feature("Configurable Lead & Lag Offsets", "Allows positive buffer delays (+days) or accelerated negative lead times (-days) between linked tasks.")
    add_bullet_feature("Milestone Diamond Markers (◆)", "Zero-duration stage gates, regulatory submissions, and key customer deliverables displayed as distinctive diamond milestones.")
    add_bullet_feature("Multi-Scale Timeline Zoom", "Instant toggling between Day, Week, and Month zoom resolutions with dynamic month/week timeline header bars.")

    # ----------------------------------------------------
    # MODULE 2: WBS TREE & PROGRESS ROLLUPS
    # ----------------------------------------------------
    add_section_header(2, "Work Breakdown Structure (WBS Tree & Rollups)")
    doc.add_paragraph("Decomposes complex engineering programs into manageable hierarchical deliverables with automatic recursive aggregation of financial costs and physical completion percentages.")

    add_bullet_feature("Hierarchical Code Generation", "Automated multi-level recursive numbering hierarchy (e.g. 1.0 Vehicle Architecture -> 1.1 Powertrain -> 1.1.1 800V SiC Inverter).")
    add_bullet_feature("Automated Cost Rollups", "Child work package budgets and actual expenditures automatically aggregate up through intermediate deliverables to top-level phases.")
    add_bullet_feature("Weighted Physical Progress Rollup", "Phase-level completion percentages calculate dynamically based on the completion state of underlying child tasks.")
    add_bullet_feature("Interactive Tree Expansion", "One-click 'Expand All' and 'Collapse All' tree controls with visual indentation levels and progress bars.")
    add_bullet_feature("Deliverable Classification", "Categorization of nodes into Phases, Major Deliverables, Work Packages, or Execution Tasks.")

    # ----------------------------------------------------
    # MODULE 3: RESOURCE CAPACITY & EQUIPMENT MANAGEMENT
    # ----------------------------------------------------
    add_section_header(3, "Resource Capacity & Specialized Equipment Planning")
    doc.add_paragraph("Manages human engineering capacity against standard weekly workloads while scheduling high-value hardware test rigs, dyno benches, and wind tunnels.")

    add_bullet_feature("40h/Week Engineer Capacity Matrix", "Tracks weekly workload allocations per engineer against standard 40-hour industrial capacity thresholds.")
    add_bullet_feature("Automated Overload Bottleneck Alerts", "Visual badges flag capacity states: Red for Overloaded (>100%), Amber for Balanced (70-100%), and Green for Available (<70%).")
    add_bullet_feature("Engineering Skills & Competency Matrix", "Tags engineering proficiencies (e.g. 800V SiC Architecture, Immersion Cooling CFD, Solid-State LiDAR, ISO 26262 ASIL-D Audit).")
    add_bullet_feature("Hourly Cost Rate Modeling", "Calculates burn rates and labor costs based on specialized hourly financial profiles per engineer.")
    add_bullet_feature("Hardware Test Rigs & Equipment Ledger", "Schedules specialized capital equipment (e.g. 500kW Dual-Motor Powertrain Dyno, Aeroacoustic Wind Tunnel, Immersion Safety Chamber).")

    # ----------------------------------------------------
    # MODULE 4: FINANCIALS & EARNED VALUE MANAGEMENT (EVM)
    # ----------------------------------------------------
    add_section_header(4, "Financial Controls & Earned Value Management (EVM)")
    doc.add_paragraph("Industrial-standard EVM performance measurement baseline for monitoring cost efficiency, schedule progress, and forecasting project completion financials.")

    add_bullet_feature("Core EVM Financial Metrics", "Real-time calculation of Budget at Completion (BAC), Planned Value (PV), Earned Value (EV), and Actual Cost (AC).")
    add_bullet_feature("Cost Performance Index (CPI = EV / AC)", "Calculates cost efficiency gauge (Green > 1.0 indicates budget surplus; Red < 1.0 indicates cost overrun).")
    add_bullet_feature("Schedule Performance Index (SPI = EV / PV)", "Calculates schedule efficiency gauge (Green > 1.0 indicates ahead of schedule; Amber < 1.0 indicates schedule delay).")
    add_bullet_feature("Estimate at Completion (EAC = BAC / CPI)", "Forecasts final total project expenditure based on observed historical performance.")
    add_bullet_feature("Variance at Completion (VAC = BAC - EAC)", "Projects expected budget variance surplus or overrun at program closeout.")
    add_bullet_feature("Historical EVM S-Curve Visualization", "Interactive Recharts multi-line trend graph comparing Planned Value, Earned Value, and Actual Cost trajectories over time.")
    add_bullet_feature("Cost Breakdown Structure (CBS) Ledger", "Tracks purchase orders (POs), external supplier contracts, tooling invoices, and component reconciliations.")

    # ----------------------------------------------------
    # MODULE 5: STAGE-GATE GOVERNANCE (APQP G0 - G4)
    # ----------------------------------------------------
    add_section_header(5, "Stage-Gate Governance & Quality Gatekeeper (APQP G0 to G4)")
    doc.add_paragraph("Structured automotive APQP (Advanced Product Quality Planning) stage-gate governance framework ensuring strict quality compliance prior to phase advancement.")

    add_bullet_feature("5 Standardized Automotive Phase Gates", "Includes G0 (Concept Feasibility), G1 (Detailed Design Freeze), G2 (Tooling Release), G3 (Homologation & Crash Test Pass), and G4 (SOP Series Assembly).")
    add_bullet_feature("Entry & Exit Criteria Checklists", "Mandatory quality checklists with real-time pass/fail toggle verification.")
    add_bullet_feature("Executive Gate Sign-Off Modal", "Audit-trailed approval workflows capturing Sign-Off Decision (APPROVED, CONDITIONAL_PASS, REJECTED), Signatory Name, Timestamp, and Steering Committee notes.")
    add_bullet_feature("Readiness Score Indicators", "Visual percentage progress indicators reflecting completed prerequisite criteria per gate.")

    # ----------------------------------------------------
    # MODULE 6: ENGINEERING CHANGE MANAGEMENT (ECO / CCB)
    # ----------------------------------------------------
    add_section_header(6, "Engineering Change Orders (ECO) & Change Control Board (CCB)")
    doc.add_paragraph("Rigorous change management workflow evaluating the triple-constraint impact of design modifications before formal Change Control Board approval.")

    add_bullet_feature("Engineering Change Request Ledger", "Tracks all formal modifications (e.g. ECO-0104 Silicon Carbide MOSFET upgrade, ECO-0105 Dielectric immersion fluid change).")
    add_bullet_feature("Triple-Constraint Impact Analysis", "Quantifies financial delta (ΔCost in Lakhs), schedule impact (ΔSchedule in Days), and technical risk classification.")
    add_bullet_feature("Formal CCB Voting Actions", "In-line executive actions to Approve ECO, Reject Request, or Mark Changes as Formally Implemented.")
    add_bullet_feature("Audit Trail & Version Traceability", "Maintains complete record of requesting engineer, reason for change, and approving authority.")

    # ----------------------------------------------------
    # MODULE 7: RISK MANAGEMENT & 5x5 HEATMAP
    # ----------------------------------------------------
    add_section_header(7, "Risk Management & 5x5 Probability-Impact Heatmap")
    doc.add_paragraph("Visualizes program risks, quantifies exposure scores, and assigns proactive mitigation and contingency remediation plans.")

    add_bullet_feature("5x5 Probability-Impact Heatmap Grid", "Interactive visual risk matrix color-coded by severity (Crimson Extreme, Orange High, Yellow Moderate, Green Low).")
    add_bullet_feature("Quantitative Risk Exposure Score", "Calculated as Risk Score = Probability (1 to 5) x Impact (1 to 5).")
    add_bullet_feature("Proactive Mitigation & Contingency Strategy", "Structured fields for root cause triggers, mitigation actions, fallback plans, and risk owners.")
    add_bullet_feature("Interactive Matrix Cell Filtering", "Clicking any cell in the 5x5 grid filters the risk ledger to display only matching items.")

    # ----------------------------------------------------
    # MODULE 8: ISSUE MANAGEMENT & CAPA (RCA)
    # ----------------------------------------------------
    add_section_header(8, "Issue Management & Corrective Action (CAPA / RCA)")
    doc.add_paragraph("Tracks active engineering defects, non-conformances, and field anomalies with systematic Root Cause Analysis (RCA) and Corrective & Preventive Actions (CAPA).")

    add_bullet_feature("Severity-Based Issue Categorization", "Categorizes blocking issues into Critical, High, Medium, and Low severity tiers.")
    add_bullet_feature("Root Cause Analysis (RCA) Documentation", "Captures detailed underlying systemic failure modes and physical causes.")
    add_bullet_feature("CAPA Remediation Tracking", "Logs corrective actions, preventative measures, assigned owners, and target resolution dates.")
    add_bullet_feature("Lifecycle Status Progression", "Manages issues through OPEN, IN_INVESTIGATION, CAPA_PENDING, and RESOLVED states.")

    # ----------------------------------------------------
    # MODULE 9: DECISION LOG & INSTITUTIONAL MEMORY
    # ----------------------------------------------------
    add_section_header(9, "Decision Log & Institutional Memory")
    doc.add_paragraph("Maintains an immutable historical record of major technical trade-offs, architecture selections, and executive steering committee rulings.")

    add_bullet_feature("Architectural Decision Registry", "Logs formal decisions (e.g. DEC-001 Selection of 800V SiC over 400V IGBT, DEC-002 Direct Immersion Cooling).")
    add_bullet_feature("Impact Quantifier Badges", "Displays financial delta (ΔCost) and schedule variance (ΔWeeks) associated with each architectural choice.")
    add_bullet_feature("Executive Approver Signatures", "Links decisions directly to authorizing program directors and technical leads.")

    # ----------------------------------------------------
    # MODULE 10: BASELINE SCHEDULE & SLIPPAGE ANALYSIS
    # ----------------------------------------------------
    add_section_header(10, "Schedule Baseline (T0) & Slippage Analysis")
    doc.add_paragraph("Locks approved project timelines and tracks real-time schedule drift against the authorized baseline.")

    add_bullet_feature("T0 Baseline Snapshot Freeze", "Creates an immutable snapshot of task planned start and due dates.")
    add_bullet_feature("Schedule Variance Slippage (ΔDays)", "Automatically calculates variance between current scheduled dates and locked baseline dates.")
    add_bullet_feature("Historical Baseline Revisions", "Supports managing multiple baseline revisions with date stamps and authorization notes.")

    # ----------------------------------------------------
    # MODULE 11: AGILE EXECUTION, COLLABORATION & SECURITY
    # ----------------------------------------------------
    add_section_header(11, "Agile Execution, Collaboration & Enterprise Security")
    doc.add_paragraph("Combines high-level project controls with granular day-to-day execution tools and security infrastructure.")

    add_bullet_feature("Interactive Kanban Board", "Drag-and-drop task movement across Concept, Simulation, Prototype, Track Testing, and Homologation columns.")
    add_bullet_feature("Subtask Checklist Progress", "Interactive checklists with automatic percentage completion progress bars.")
    add_bullet_feature("Real-Time WebSockets (Socket.io)", "Instant two-way synchronization across all active browser sessions.")
    add_bullet_feature("Role-Based Access Control (RBAC)", "Multi-tiered permissions supporting ADMIN, MANAGER, MEMBER, and VIEWER roles.")
    add_bullet_feature("Live Time Tracking Widget", "Integrated stopwatch timer allowing engineers to log billable hours directly against specific work packages.")
    add_bullet_feature("FullCalendar Multi-View Timeline", "Integrated calendar supporting Month, Week, Day, and List views with milestone diamond overlays.")
    add_bullet_feature("User Profile & Photo Customizer", "Interactive profile settings supporting 1-click avatar presets, custom image URLs, and password encryption.")

    # ----------------------------------------------------
    # PRE-CONFIGURED TEAM DIRECTORY TABLE
    # ----------------------------------------------------
    add_section_header(12, "Pre-Configured Engineering Leadership Roster")
    doc.add_paragraph("The platform includes a fully seeded, cross-functional automotive engineering leadership team with dedicated roles, access privileges, and specialized skills:")

    table_team = doc.add_table(rows=7, cols=4)
    table_team.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_team.autofit = False

    headers = ["Team Member", "Role & Title", "Department", "Login Email"]
    col_widths = [Inches(1.7), Inches(2.2), Inches(1.8), Inches(1.5)]

    for j, h in enumerate(headers):
        cell = table_team.rows[0].cells[j]
        cell.width = col_widths[j]
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, 120, 120, 100, 100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    team_members = [
        ("Parth Singh (Admin)", "Program Director & Chief EV Architect", "Executive Engineering", "admin@proflow.com"),
        ("Kashish Motwani", "Battery Systems & Thermal Lead", "Energy Storage Systems", "kashish.motwani@proflow.com"),
        ("Abhinav Kumar", "Manufacturing & Tooling Specialist", "Advanced Manufacturing", "abhinav.kumar@proflow.com"),
        ("Rashika Sarawgi", "ADAS Perception & Autonomy Lead", "Autonomous Driving & SW", "rashika.sarawgi@proflow.com"),
        ("Ritik Raj", "800V SiC Powertrain & Inverter Lead", "Powertrain & Power Electronics", "ritik.raj@proflow.com"),
        ("Priyanshi Mandiya", "Quality & Homologation Gatekeeper", "Quality & Compliance", "priyanshi.mandiya@proflow.com")
    ]

    for i, row_data in enumerate(team_members):
        row = table_team.rows[i + 1]
        bg_color = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.width = col_widths[j]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, 100, 100, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            if j == 0:
                r.font.bold = True
            r.font.color.rgb = slate_dark

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Footer note
    p_foot = doc.add_paragraph()
    r_foot = p_foot.add_run("© 2026 ProFlow Project Management Suite. All rights reserved. Generated automatically via Antigravity AI Engine.")
    r_foot.font.size = Pt(8.5)
    r_foot.font.italic = True
    r_foot.font.color.rgb = slate_gray
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_path = os.path.abspath("ProFlow_Enterprise_Project_Management_Features.docx")
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_document()
